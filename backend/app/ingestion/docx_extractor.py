# backend/app/ingest/docx_extractor.py
# DVMELTSS-FIX: V - Validate, E - Error handling, S - Security, A - Async
# BATMAN-FIX: A - True async, M - Memory safety
# OWASP-FIX: 7 - PII redaction, 9 - File handling
# ✅ FIXED: Proper async/sync bridge + input validation + safe file cleanup

from __future__ import annotations

import asyncio
import logging
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Final, Optional, Any

from langchain_core.documents import Document

# DVMELTSS-M: Import centralized utilities
from app.core.ingest_utils import redact_pii, generate_ingest_correlation_id
from app.core.celery_utils import run_async_in_task  # ✅ NEW: For safe async execution

logger = logging.getLogger(__name__)

_MAX_DOC_SIZE_MB: Final = 50
_MAX_PARAGRAPHS: Final = 50000
_MAX_TABLE_CELLS: Final = 100000

# ✅ NEW: Timeout for python-docx operations (seconds)
_DOCX_TIMEOUT: Final = 120


@dataclass  # FIXED: Removed frozen=True for mutability if needed
class DocxContent:
    """Extracted content from a .docx file."""

    source_file: str
    full_text: str
    paragraphs: list[str]
    tables: list[list[list[str]]]
    headings: list[tuple[int, str]]
    metadata: dict = field(default_factory=dict)
    page_count: int = 1
    error: Optional[str] = None
    correlation_id: Optional[str] = None

    def to_dict(self) -> dict:
        """Serialize for API responses / logging (with metadata safety)."""
        # ✅ FIXED: Safe serialization with None handling
        safe_meta = {k: redact_pii(v) if isinstance(v, str) else v for k, v in (self.metadata or {}).items()}
        return {
            "source_file": self.source_file,
            "full_text_length": len(self.full_text),
            "paragraph_count": len(self.paragraphs or []),
            "table_count": len(self.tables or []),
            "heading_count": len(self.headings or []),
            "metadata": safe_meta,
            "error": self.error,
            "correlation_id": self.correlation_id,
        }


# ✅ NEW: Input validation helper
def _validate_docx_inputs(
    file_path: Optional[str | Path],
    content: Optional[DocxContent],
    correlation_id: Optional[str],
    corr_id: str,
) -> tuple[bool, str]:
    """Validate DOCX extractor inputs before processing."""
    if file_path is not None and not isinstance(file_path, (str, Path)):
        return False, "file_path must be a string, Path, or None"
    if content is not None and not isinstance(content, DocxContent):
        return False, "content must be a DocxContent instance or None"
    if correlation_id is not None and not isinstance(correlation_id, str):
        return False, "correlation_id must be a string or None"
    return True, ""


class DocxExtractor:
    """Extracts structured content from .docx Word files."""

    def __init__(self, chunk_size: int = 1024):
        self.chunk_size = chunk_size
        logger.info(f"DocxExtractor initialized: chunk_size={chunk_size}")

    def _sanitize_metadata(self, cp) -> dict:
        """OWASP-7: Extract and sanitize core properties."""
        try:
            meta = {
                "title": cp.title or "",
                "author": cp.author or "",
                "created": str(cp.created) if cp.created else "",
                "modified": str(cp.modified) if cp.modified else "",
                "subject": cp.subject or "",
            }
            for k, v in meta.items():
                if isinstance(v, str) and v:
                    meta[k] = redact_pii(v)
            return meta
        except Exception:
            return {}

    async def extract_async(
        self,
        file_path: str | Path,
        correlation_id: Optional[str] = None,
    ) -> DocxContent:
        """Async version: Extract all content from a .docx file."""
        corr_id = correlation_id or generate_ingest_correlation_id("docx")

        # ✅ Validate inputs
        is_valid, error = _validate_docx_inputs(file_path, None, correlation_id, corr_id)
        if not is_valid:
            logger.error(f"[{corr_id}] Invalid DOCX inputs: {error}")
            return DocxContent(
                source_file=str(file_path) if file_path else "unknown",
                full_text="",
                paragraphs=[],
                tables=[],
                headings=[],
                error=error,
                correlation_id=corr_id,
            )

        file_path_obj = Path(file_path) if isinstance(file_path, str) else file_path

        if not file_path_obj.exists():
            return DocxContent(
                source_file=str(file_path_obj),
                full_text="",
                paragraphs=[],
                tables=[],
                headings=[],
                error=f"File not found: {file_path_obj}",
                correlation_id=corr_id,
            )

        size_mb = file_path_obj.stat().st_size / 1024 / 1024
        if size_mb > _MAX_DOC_SIZE_MB:
            logger.warning(f"[{corr_id}] File too large: {size_mb:.1f}MB > {_MAX_DOC_SIZE_MB}MB")

        try:
            import docx
        except ImportError:
            return DocxContent(
                source_file=str(file_path_obj),
                full_text="",
                paragraphs=[],
                tables=[],
                headings=[],
                error="python-docx not installed",
                correlation_id=corr_id,
            )

        doc = None
        try:
            loop = asyncio.get_running_loop()

            # ✅ FIXED: Use wait_for with timeout for docx operations
            doc = await asyncio.wait_for(
                loop.run_in_executor(None, lambda: docx.Document(str(file_path_obj))),
                timeout=_DOCX_TIMEOUT,
            )
        except asyncio.TimeoutError:
            logger.error(f"[{corr_id}] DOCX file open timed out after {_DOCX_TIMEOUT}s")
            return DocxContent(
                source_file=str(file_path_obj),
                full_text="",
                paragraphs=[],
                tables=[],
                headings=[],
                error=f"Timeout opening DOCX file after {_DOCX_TIMEOUT}s",
                correlation_id=corr_id,
            )
        except Exception as e:
            return DocxContent(
                source_file=str(file_path_obj),
                full_text="",
                paragraphs=[],
                tables=[],
                headings=[],
                error=f"Failed to open docx: {type(e).__name__}",
                correlation_id=corr_id,
            )

        paragraphs = []
        headings = []
        all_text = []

        try:
            for i, para in enumerate(doc.paragraphs):
                if i >= _MAX_PARAGRAPHS:
                    logger.warning(f"[{corr_id}] Paragraph limit reached ({_MAX_PARAGRAPHS})")
                    break
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name.lower() if para.style else ""
                if "heading" in style_name:
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    headings.append((level, text))
                    formatted = "#" * level + " " + text
                    all_text.append(formatted)
                else:
                    paragraphs.append(text)
                    all_text.append(text)

            tables = []
            for tbl_idx, tbl in enumerate(doc.tables):
                table_data = []
                cell_count = 0
                for row in tbl.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_count += 1
                        if cell_count > _MAX_TABLE_CELLS:
                            break
                        row_data.append(cell.text.strip())
                    if cell_count > _MAX_TABLE_CELLS:
                        break
                    if any(c for c in row_data):
                        table_data.append(row_data)
                if table_data and cell_count <= _MAX_TABLE_CELLS:
                    tables.append(table_data)
                elif cell_count > _MAX_TABLE_CELLS:
                    logger.warning(f"[{corr_id}] Table {tbl_idx} exceeded cell limit")

            meta = self._sanitize_metadata(doc.core_properties)
            full_text = "\n\n".join(all_text)

            logger.info(
                f"[{corr_id}] DocxExtractor: {file_path_obj.name} | "
                f"{len(paragraphs)} paras | {len(tables)} tables | {len(headings)} headings"
            )

            return DocxContent(
                source_file=str(file_path_obj),
                full_text=full_text,
                paragraphs=paragraphs,
                tables=tables,
                headings=headings,
                metadata=meta,
                correlation_id=corr_id,
            )
        except Exception as e:
            logger.error(f"[{corr_id}] Docx extraction failed during processing: {e}")
            return DocxContent(
                source_file=str(file_path_obj),
                full_text="",
                paragraphs=[],
                tables=[],
                headings=[],
                error=f"Processing error: {type(e).__name__}",
                correlation_id=corr_id,
            )
        finally:
            # ✅ FIXED: Proper cleanup of docx document
            if doc is not None:
                try:
                    # python-docx doesn't have explicit close, but ensure no references held
                    del doc
                except Exception:
                    pass

    async def to_langchain_documents_async(
        self,
        content: DocxContent,
        correlation_id: Optional[str] = None,
    ) -> list[Document]:
        """Async wrapper for document conversion."""
        corr_id = correlation_id or content.correlation_id or generate_ingest_correlation_id("docx_chunks")

        # ✅ Validate inputs
        is_valid, error = _validate_docx_inputs(None, content, correlation_id, corr_id)
        if not is_valid or not content.full_text:
            logger.warning(f"Invalid content or empty text: {error}")
            return []

        chunks = []
        current = []
        word_count = 0

        # ✅ FIXED: Safe iteration over paragraphs
        for para in content.paragraphs or []:
            if not isinstance(para, str) or not para.strip():
                continue
            words = len(para.split())
            if word_count + words > self.chunk_size and current:
                chunk_text = "\n\n".join(current)
                if chunk_text.strip():
                    chunks.append(chunk_text)
                current = []
                word_count = 0
            current.append(para)
            word_count += words

        if current:
            chunk_text = "\n\n".join(current)
            if chunk_text.strip():
                chunks.append(chunk_text)

        # ✅ FIXED: Safe iteration over tables
        for i, table in enumerate(content.tables or []):
            if not isinstance(table, list):
                continue
            table_text = "\n".join(
                " | ".join(row)
                for row in table
                if isinstance(row, list) and any(c.strip() for c in row if isinstance(c, str))
            )
            if table_text.strip():
                chunks.append(f"[Table {i+1}]\n{table_text}")

        now = datetime.now(timezone.utc).isoformat()
        docs = []

        for i, chunk in enumerate(chunks):
            if not chunk or not chunk.strip():
                continue
            docs.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source_file": content.source_file,
                        "page_number": 0,
                        "chunk_id": str(uuid.uuid4()),
                        "parent_id": "",
                        "block_type": "paragraph",
                        "language": "en",
                        "ocr_confidence": 1.0,
                        "chunk_type": "child",
                        "ingest_timestamp": now,
                        "document_type": "docx",
                        "char_count": len(chunk),
                        "docx_title": content.metadata.get("title", ""),
                        "docx_author": content.metadata.get("author", ""),
                        "correlation_id": corr_id,
                    },
                )
            )
        return docs

    def extract(
        self,
        file_path: str | Path,
        correlation_id: Optional[str] = None,
    ) -> DocxContent:
        """
        Sync wrapper — prefers async version in new code.
        ✅ FIXED: Use run_async_in_task helper to avoid deadlock.
        """

        async def _do_extract():
            return await self.extract_async(file_path, correlation_id)

        return run_async_in_task(_do_extract)

    def to_langchain_documents(
        self,
        content: DocxContent,
        chunk_size: int = 1024,
        correlation_id: Optional[str] = None,
    ) -> list[Document]:
        """
        Sync wrapper — prefers async version in new code.
        ✅ FIXED: Use run_async_in_task helper to avoid deadlock.
        """

        async def _do_convert():
            old_size = self.chunk_size
            self.chunk_size = chunk_size
            try:
                return await self.to_langchain_documents_async(content, correlation_id)
            finally:
                self.chunk_size = old_size

        return run_async_in_task(_do_convert)


def get_docx_metadata() -> dict[str, Any]:
    """✅ NEW: Return DOCX extractor metadata for debugging."""
    return {
        "limits": {
            "max_doc_size_mb": _MAX_DOC_SIZE_MB,
            "max_paragraphs": _MAX_PARAGRAPHS,
            "max_table_cells": _MAX_TABLE_CELLS,
            "default_chunk_size": 1024,
        },
        "docx_timeout_seconds": _DOCX_TIMEOUT,
        "pii_redaction_enabled": True,
        "metadata_sanitization_enabled": True,
    }


# DVMELTSS-M: Explicit module exports
__all__ = [
    "DocxExtractor",
    "DocxContent",
    "get_docx_metadata",
]
# Local smoke test entry point. Run: python -m
if __name__ == "__main__":
    import sys
    from app.core.module_smoke import run_module_smoke

    run_module_smoke(sys.modules[__name__], __file__)
